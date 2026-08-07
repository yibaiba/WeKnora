from io import BytesIO
import unittest

from docx import Document as WordDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

from function_test_loader import load_function_tests

from docreader.models.document import Document
from docreader.parser.docx_structure import attach_docx_structure
from docreader.structure import structure_blocks_from_document


def _png_bytes() -> bytes:
    output = BytesIO()
    Image.new("RGB", (12, 8), "white").save(output, format="PNG")
    return output.getvalue()


def _structured_docx() -> bytes:
    document = WordDocument()
    document.add_heading("Guide", level=1)
    outline_heading = document.add_paragraph("Outline section")
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "1")
    outline_heading._p.get_or_add_pPr().append(outline)
    document.add_paragraph("First item", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "State"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "Ready"
    document.add_picture(BytesIO(_png_bytes()))
    document.add_paragraph("Figure 1: Overview", style="Caption")
    page_paragraph = document.add_paragraph("Second page")
    page_paragraph.paragraph_format.page_break_before = True
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _final_markdown() -> str:
    return (
        "# Guide\n\n"
        "## Outline section\n\n"
        "- First item\n\n"
        "| Name | State |\n"
        "| --- | --- |\n"
        "| Alpha | Ready |\n\n"
        "![image](images/picture.png)\n"
        "Figure 1: Overview\n\n"
        "Second page\n"
    )


def test_docx_native_structure_maps_final_markdown_and_transport():
    markdown = _final_markdown()
    document = attach_docx_structure(_structured_docx(), Document(content=markdown))

    kinds = [block.kind for block in document.structure_blocks]
    assert kinds == [
        "heading",
        "heading",
        "list_item",
        "table_header",
        "table_row",
        "image_caption",
        "page_region",
    ]
    assert all(markdown[block.start : block.end] for block in document.structure_blocks)
    assert len({block.id for block in document.structure_blocks}) == len(
        document.structure_blocks
    )
    heading = document.structure_blocks[0]
    section = document.structure_blocks[1]
    assert section.parent_id == heading.id
    assert all(
        block.parent_id == section.id for block in document.structure_blocks[2:]
    )
    header = document.structure_blocks[3]
    row = document.structure_blocks[4]
    assert header.table_id == row.table_id
    assert "| --- | --- |" in markdown[header.start : header.end]
    image = document.structure_blocks[5]
    assert image.context_kinds == ["image", "caption"]

    transported = structure_blocks_from_document(document)
    assert [block.kind for block in transported] == kinds
    assert transported[3].table_id == header.table_id


def test_markitdown_docx_production_path_returns_source_aligned_blocks():
    try:
        from docreader.parser.markitdown_parser import MarkitdownParser
    except ImportError as error:
        raise unittest.SkipTest("MarkItDown runtime dependencies unavailable") from error

    document = MarkitdownParser(
        file_name="semantic-structure.docx", file_type="docx"
    ).parse_into_text(_structured_docx())

    assert document.content
    assert {"heading", "list_item", "table_header", "table_row"}.issubset(
        {block.kind for block in document.structure_blocks}
    )
    assert all(
        document.content[block.start : block.end]
        for block in document.structure_blocks
    )


def test_damaged_outline_metadata_does_not_forge_structure_or_change_markdown():
    source = WordDocument()
    paragraph = source.add_paragraph("Body text")
    properties = paragraph._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "invalid")
    properties.append(outline)
    output = BytesIO()
    source.save(output)
    markdown = "Body text\n"

    document = attach_docx_structure(output.getvalue(), Document(content=markdown))

    assert document.content == markdown
    assert document.structure_blocks == []


def test_invalid_docx_metadata_exposes_reason_without_body_content():
    markdown = "Preserved body"

    document = attach_docx_structure(b"not-a-docx", Document(content=markdown))

    assert document.content == markdown
    assert document.structure_blocks == []
    assert document.metadata == {"semantic_structure_status": "docx_metadata_invalid"}
    assert markdown not in document.metadata["semantic_structure_status"]


def load_tests(loader, tests, pattern):
    del loader, pattern
    return load_function_tests(globals(), tests)
