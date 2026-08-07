"""DOCX-native structural hints mapped onto finalized Markdown."""

from __future__ import annotations

import html
import logging
import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from typing import Iterable

from docx import Document as WordDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph

from docreader.models.document import Document, StructureBlock


_HEADING_STYLE = re.compile(r"^(?:heading|标题)\s*([1-6])$", re.IGNORECASE)
_CAPTION_STYLE = re.compile(r"^(?:caption|题注)$", re.IGNORECASE)
_CAPTION_TEXT = re.compile(r"^(?:fig(?:ure)?\.?|图(?:片)?|插图)\s*\d+", re.IGNORECASE)
_LIST_STYLE = re.compile(r"(?:list|列表)", re.IGNORECASE)
_MARKDOWN_PREFIX = re.compile(r"^(?:#{1,6}\s+|[-*+]\s+|\d+[.)]\s+|>\s*)")
_IMAGE_LINE = re.compile(r"^\s*!\[[^]]*]\([^)]+\)\s*$")
_TABLE_SEPARATOR = re.compile(r"^:?-{3,}:?$")

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class _MarkdownLine:
    text: str
    start: int
    end: int


class _DocxStructureMapper:
    def __init__(self, markdown: str):
        self.lines = _markdown_lines(markdown)
        self.cursor = 0
        self.blocks: list[StructureBlock] = []
        self.heading_stack: dict[int, str] = {}
        self.sequences: dict[str, int] = {}

    def map(self, elements: Iterable[Paragraph | Table]) -> list[StructureBlock]:
        for element in elements:
            if isinstance(element, Paragraph):
                self._map_paragraph(element)
            elif isinstance(element, Table):
                self._map_table(element)
        return self.blocks

    def _map_paragraph(self, paragraph: Paragraph) -> None:
        text = paragraph.text.strip()
        if not text:
            return
        line_index = self._find_text_line(text)
        if line_index is None:
            return
        line = self.lines[line_index]
        depth = _heading_depth(paragraph)
        if depth is not None:
            self._add_heading(line, depth)
        elif _is_list(paragraph):
            self._add_block("list_item", line.start, line.end, atomic=True)
        elif _is_caption(paragraph, text):
            self._add_caption(line_index)
        elif _has_page_break(paragraph):
            self._add_block("page_region", line.start, line.end, atomic=False)
        self.cursor = max(self.cursor, line.end)

    def _map_table(self, table: Table) -> None:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if not rows:
            return
        line_indexes = self._find_table_lines(rows)
        if not line_indexes:
            return
        table_id = self._next_id("table")
        header_index, separator_index = line_indexes[0], line_indexes[1]
        header = self.lines[header_index]
        separator = self.lines[separator_index]
        self._add_block(
            "table_header", header.start, separator.end,
            atomic=True, table_id=table_id,
        )
        for line_index in line_indexes[2:]:
            line = self.lines[line_index]
            self._add_block(
                "table_row", line.start, line.end,
                atomic=True, table_id=table_id,
            )
        self.cursor = max(self.cursor, self.lines[line_indexes[-1]].end)

    def _find_text_line(self, source_text: str) -> int | None:
        expected = _normalize_text(source_text)
        for index, line in enumerate(self.lines):
            if line.start < self.cursor or _is_table_line(line.text):
                continue
            if _normalize_markdown_line(line.text) == expected:
                return index
        return None

    def _find_table_lines(self, rows: list[list[str]]) -> list[int]:
        expected = [[_normalize_text(cell) for cell in row] for row in rows]
        for index, line in enumerate(self.lines):
            if line.start < self.cursor or _markdown_cells(line.text) != expected[0]:
                continue
            if index + 1 >= len(self.lines) or not _is_table_separator(self.lines[index + 1].text):
                continue
            matches = [index, index + 1]
            next_index = index + 2
            for row in expected[1:]:
                if next_index >= len(self.lines) or _markdown_cells(self.lines[next_index].text) != row:
                    break
                matches.append(next_index)
                next_index += 1
            if len(matches) == len(expected) + 1:
                return matches
        return []

    def _add_heading(self, line: _MarkdownLine, depth: int) -> None:
        block_id = self._next_id("heading")
        parent_id = _nearest_heading(self.heading_stack, depth - 1)
        self._add_block(
            "heading", line.start, line.end, block_id=block_id,
            parent_id=parent_id, section_depth=depth, atomic=True,
        )
        self.heading_stack = {
            level: value for level, value in self.heading_stack.items() if level < depth
        }
        self.heading_stack[depth] = block_id

    def _add_caption(self, line_index: int) -> None:
        image_index = line_index - 1
        while image_index >= 0 and not self.lines[image_index].text.strip():
            image_index -= 1
        if image_index < 0 or not _IMAGE_LINE.match(self.lines[image_index].text):
            return
        self._add_block(
            "image_caption", self.lines[image_index].start, self.lines[line_index].end,
            atomic=True, context_kinds=["image", "caption"],
        )

    def _add_block(
        self,
        kind: str,
        start: int,
        end: int,
        *,
        atomic: bool,
        block_id: str = "",
        parent_id: str | None = None,
        section_depth: int = 0,
        table_id: str = "",
        context_kinds: list[str] | None = None,
    ) -> None:
        self.blocks.append(StructureBlock(
            id=block_id or self._next_id(kind), kind=kind, start=start, end=end,
            parent_id=(
                _nearest_heading(self.heading_stack, 6)
                if parent_id is None else parent_id
            ),
            section_depth=section_depth, table_id=table_id, atomic=atomic,
            confidence="high", context_kinds=context_kinds or [],
        ))

    def _next_id(self, kind: str) -> str:
        self.sequences[kind] = self.sequences.get(kind, 0) + 1
        return f"docx-{kind}-{self.sequences[kind]:04d}"


def attach_docx_structure(content: bytes, document: Document) -> Document:
    """Return the parsed document with OOXML-backed, final-Markdown hints."""
    try:
        word_document = WordDocument(BytesIO(content))
    except (PackageNotFoundError, ValueError, KeyError, zipfile.BadZipFile) as exc:
        logger.warning("DOCX structure metadata is invalid: %s", exc)
        metadata = dict(document.metadata)
        metadata["semantic_structure_status"] = "docx_metadata_invalid"
        return document.model_copy(update={"metadata": metadata})
    blocks = _DocxStructureMapper(document.content).map(word_document.iter_inner_content())
    return document.model_copy(update={"structure_blocks": blocks})


def _markdown_lines(markdown: str) -> list[_MarkdownLine]:
    result: list[_MarkdownLine] = []
    offset = 0
    for value in markdown.splitlines(keepends=True):
        result.append(_MarkdownLine(text=value, start=offset, end=offset + len(value)))
        offset += len(value)
    if offset < len(markdown):
        result.append(_MarkdownLine(markdown[offset:], offset, len(markdown)))
    return result


def _normalize_text(value: str) -> str:
    value = html.unescape(value).replace("\\|", "|")
    return " ".join(value.split())


def _normalize_markdown_line(value: str) -> str:
    value = _MARKDOWN_PREFIX.sub("", value.strip())
    value = re.sub(r"([*_~`])", "", value)
    value = re.sub(r"\\([\\`*{}\[\]()#+.!_|>-])", r"\1", value)
    return _normalize_text(value)


def _markdown_cells(value: str) -> list[str]:
    stripped = value.strip()
    if not stripped.startswith("|"):
        return []
    cells = re.split(r"(?<!\\)\|", stripped.strip("|"))
    return [_normalize_markdown_line(cell) for cell in cells]


def _is_table_line(value: str) -> bool:
    return bool(_markdown_cells(value))


def _is_table_separator(value: str) -> bool:
    cells = _markdown_cells(value)
    return bool(cells) and all(_TABLE_SEPARATOR.fullmatch(cell) for cell in cells)


def _heading_depth(paragraph: Paragraph) -> int | None:
    style = _paragraph_style(paragraph)
    style_name = style.name.strip() if style else ""
    match = _HEADING_STYLE.match(style_name)
    if match:
        return int(match.group(1))
    for properties in (paragraph._p.pPr, getattr(getattr(style, "element", None), "pPr", None)):
        outline = getattr(properties, "outlineLvl", None)
        try:
            if outline is not None:
                return max(1, min(6, int(outline.val) + 1))
        except (TypeError, ValueError):
            return None
    return None


def _is_list(paragraph: Paragraph) -> bool:
    properties = paragraph._p.pPr
    if properties is not None and properties.numPr is not None:
        return True
    style = _paragraph_style(paragraph)
    style_name = style.name if style else ""
    return bool(_LIST_STYLE.search(style_name))


def _is_caption(paragraph: Paragraph, text: str) -> bool:
    style = _paragraph_style(paragraph)
    style_name = style.name.strip() if style else ""
    return bool(_CAPTION_STYLE.match(style_name) or _CAPTION_TEXT.match(text))


def _has_page_break(paragraph: Paragraph) -> bool:
    if paragraph.paragraph_format.page_break_before:
        return True
    return bool(paragraph._element.xpath(
        ".//w:lastRenderedPageBreak | .//w:br[@w:type='page']"
    ))


def _nearest_heading(stack: dict[int, str], maximum: int) -> str:
    for depth in range(maximum, 0, -1):
        if stack.get(depth):
            return stack[depth]
    return ""


def _paragraph_style(paragraph: Paragraph):
    try:
        return paragraph.style
    except (AttributeError, KeyError, ValueError):
        return None
